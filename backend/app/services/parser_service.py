import pdfplumber
import docx  # NEW - For DOCX support
import re
import io
from typing import List, Dict

# ==================== EXPANDED SKILL TAXONOMY ====================

SKILL_ALIASES = {
    # Programming Languages
    "python": ["python", "python3", "py"],
    "java": ["java"],
    "javascript": ["javascript", "js", "ecmascript"],
    "typescript": ["typescript", "ts"],
    "c++": ["c++", "cpp", "cplusplus"],
    "c#": ["c#", "csharp", "c sharp"],
    "go": ["golang", "go"],
    "rust": ["rust"],
    "ruby": ["ruby"],
    "php": ["php"],
    "swift": ["swift"],
    "kotlin": ["kotlin"],
    
    # Frontend
    "react": ["react", "reactjs", "react.js"],
    "angular": ["angular", "angularjs"],
    "vue": ["vue", "vuejs", "vue.js"],
    "next.js": ["next", "nextjs", "next.js"],
    "svelte": ["svelte"],
    "html": ["html", "html5"],
    "css": ["css", "css3"],
    "tailwind": ["tailwind", "tailwindcss"],
    "bootstrap": ["bootstrap"],
    "material-ui": ["material-ui", "mui", "material ui"],
    
    # Backend
    "node.js": ["node", "nodejs", "node.js"],
    "django": ["django"],
    "fastapi": ["fastapi", "fast api"],
    "flask": ["flask"],
    "spring": ["spring", "spring boot", "springboot"],
    "express": ["express", "expressjs", "express.js"],
    "laravel": ["laravel"],
    "asp.net": ["asp.net", "aspnet", "asp net"],
    
    # Databases
    "sql": ["sql"],
    "postgresql": ["postgresql", "postgres", "psql"],
    "mysql": ["mysql"],
    "mongodb": ["mongodb", "mongo"],
    "redis": ["redis"],
    "cassandra": ["cassandra"],
    "dynamodb": ["dynamodb"],
    "elasticsearch": ["elasticsearch", "elastic"],
    "oracle": ["oracle"],
    "sqlite": ["sqlite"],
    
    # DevOps & Cloud
    "docker": ["docker"],
    "kubernetes": ["kubernetes", "k8s"],
    "aws": ["aws", "amazon web services"],
    "azure": ["azure", "microsoft azure"],
    "gcp": ["gcp", "google cloud", "google cloud platform"],
    "jenkins": ["jenkins"],
    "gitlab ci": ["gitlab ci", "gitlab"],
    "github actions": ["github actions"],
    "terraform": ["terraform"],
    "ansible": ["ansible"],
    "git": ["git"],
    "linux": ["linux", "unix"],
    "nginx": ["nginx"],
    "apache": ["apache"],
    
    # ML & Data Science
    "machine learning": ["machine learning", "ml", "machine-learning"],
    "deep learning": ["deep learning", "dl"],
    "tensorflow": ["tensorflow", "tf"],
    "pytorch": ["pytorch", "torch"],
    "pandas": ["pandas"],
    "numpy": ["numpy"],
    "scikit-learn": ["scikit-learn", "sklearn", "scikit learn"],
    "keras": ["keras"],
    "data analysis": ["data analysis"],
    "tableau": ["tableau"],
    "power bi": ["power bi", "powerbi"],
    
    # APIs & Protocols
    "rest": ["rest", "restful", "rest api"],
    "graphql": ["graphql", "graph ql"],
    "grpc": ["grpc"],
    "websocket": ["websocket", "websockets"],
    
    # Tools
    "jira": ["jira"],
    "postman": ["postman"],
    "vs code": ["vs code", "vscode", "visual studio code"],
    "intellij": ["intellij"],
}

# Create flat lookup: variant -> canonical name
ALL_SKILL_VARIANTS = {}
for canonical, variants in SKILL_ALIASES.items():
    for variant in variants:
        ALL_SKILL_VARIANTS[variant.lower()] = canonical

# ==================== RESUME PARSER ====================

class ResumeParser:
    """
    Enhanced resume parser with:
    - PDF and DOCX support
    - Name, email, phone, skills extraction
    - Education and experience parsing
    - Dynamic confidence scoring
    """
    
    def parse(self, file_bytes: bytes, filename: str = "") -> Dict:
        """
        Parse PDF or DOCX resume and extract structured data
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (to detect file type)
            
        Returns:
            Dict with parsed resume data matching API schema
        """
        try:
            # Detect file type and extract text
            if filename.lower().endswith('.docx'):
                print(f"📄 Parsing DOCX file: {filename}")
                text = self._extract_text_from_docx(file_bytes)
            elif filename.lower().endswith('.pdf'):
                print(f"📄 Parsing PDF file: {filename}")
                text = self._extract_text_from_pdf(file_bytes)
            else:
                print(f"⚠️ Warning: Unknown file type: {filename}")
                return self._empty_result()
            
            if not text.strip():
                print("⚠️ Warning: Extracted empty text from file")
                return self._empty_result()
            
            print(f"✅ Extracted {len(text)} characters from file")
            
            # Extract individual fields
            name = self._extract_name(text)
            email = self._extract_email(text)
            phone = self._extract_phone(text)
            skills = self._extract_skills(text)
            education = self._extract_education(text)
            experience = self._extract_experience(text)
            projects = self._extract_projects(text)
            
            # Build result
            result = {
                "name": name or "Unknown Candidate",
                "email": email or "",
                "phone": phone or "",
                "skills": skills,
                "education": education,
                "experience": experience,
                "projects": projects,
                "confidence_scores": {
                    "name": self._name_confidence(name, text),
                    "email": 0.95 if email else 0.10,
                    "phone": self._phone_confidence(phone, text),
                    "skills": self._skills_confidence(skills),
                    "education": self._education_confidence(education),
                    "experience": self._experience_confidence(experience)
                }
            }
            
            print(f"✅ Parsed: {result['name']} | {len(skills)} skills | {len(education)} education | {len(experience)} experience")
            
            return result
            
        except Exception as e:
            print(f"❌ Error parsing file: {e}")
            import traceback
            traceback.print_exc()
            return self._empty_result()
    
    # ==================== TEXT EXTRACTION METHODS ====================
    
    def _extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract all text from PDF"""
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract all text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            
            return text
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            return ""
    
    # ==================== FIELD EXTRACTION METHODS ====================
    
    def _extract_name(self, text: str) -> str:
        """
        Extract candidate name with improved heuristics
        
        Strategy:
        1. Skip common resume headers
        2. Skip lines with email/phone (contact section)
        3. Find first capitalized line without numbers
        """
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        
        # Common headers to skip
        skip_headers = {
            'resume', 'curriculum vitae', 'cv', 'profile', 
            'personal information', 'contact', 'biodata'
        }
        
        for line in lines[:10]:  # Check first 10 lines only
            line_lower = line.lower()
            
            # Skip known headers
            if line_lower in skip_headers:
                continue
            
            # Skip lines with contact info
            if '@' in line or re.search(r'\d{10}', line):
                continue
            
            # Valid name checks
            if (len(line) < 50 and  # Not too long
                len(line.split()) <= 5 and  # Max 5 words
                line[0].isupper() and  # Starts with capital letter
                not re.search(r'\d', line)):  # No numbers
                
                return line
        
        # Fallback: return first non-empty line
        return lines[0] if lines else ""
    
    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        pattern = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'
        match = re.search(pattern, text)
        
        if match:
            email = match.group(0)
            # Basic validation
            if len(email) < 100 and '@' in email:
                return email
        
        return ""
    
    def _extract_phone(self, text: str) -> str:
        """
        Extract phone number with context awareness
        
        Prefers Indian format (+91...) found near keywords like "phone", "mobile"
        Falls back to generic 10-digit number
        """
        # Try Indian format with context
        pattern = r'(?:phone|mobile|contact|tel|cell)?\s*:?\s*[\+]?91[-\s]?[6-9]\d{9}'
        match = re.search(pattern, text, re.IGNORECASE)
        
        if match:
            # Extract just the number part
            number_match = re.search(r'[\+]?91[-\s]?[6-9]\d{9}', match.group(0))
            if number_match:
                return number_match.group(0)
        
        # Fallback: Generic 10-digit Indian number
        pattern = r'\b[6-9]\d{9}\b'
        match = re.search(pattern, text)
        
        if match:
            number = match.group(0)
            # Avoid false positives (years, IDs)
            if not (number.startswith('19') or number.startswith('20')):
                return number
        
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """
        Extract skills using alias matching
        
        Matches skill variants (e.g., "ReactJS", "React.js") to canonical names
        Returns deduplicated list of canonical skill names
        """
        text_lower = text.lower()
        found_skills = set()
        
        # Match each variant, map to canonical name
        for variant, canonical in ALL_SKILL_VARIANTS.items():
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(variant) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(canonical)
        
        return sorted(list(found_skills))
    
    def _extract_education(self, text: str) -> List[Dict]:
        """
        Extract education information
        
        Looks for:
        - Degrees (B.Tech, M.Tech, B.E, etc.)
        - Institutions (NIT, IIT, University, etc.)
        - Years (2020-2024, 2020 - Present, etc.)
        - CGPA/Percentage
        """
        education = []
        
        # Common degree patterns
        degree_pattern = r'(B\.?Tech|M\.?Tech|B\.?E\.?|M\.?E\.?|B\.?Sc|M\.?Sc|BCA|MCA|MBA|B\.?A\.?|M\.?A\.?|PhD|Ph\.?D\.?)'
        
        # Institution patterns
        institution_pattern = r'(NIT|IIT|IIIT|BITS|VIT|SRM|Amity|Manipal|University|College|Institute|School)\s+[\w\s,]+'
        
        # Year patterns (2020-2024, 2020 - 2024, 2020-Present, etc.)
        year_pattern = r'(19|20)\d{2}\s*[-–—]\s*((19|20)\d{2}|Present|present|Current|current)'
        
        # CGPA/Percentage patterns
        cgpa_pattern = r'(CGPA|GPA|cgpa|gpa)[:\s]+(\d+\.?\d*)\s*/?\s*(\d+)?'
        percentage_pattern = r'(\d{2,3}\.?\d*)\s*%'
        
        # Find all matches
        degrees = re.findall(degree_pattern, text, re.IGNORECASE)
        institutions = re.findall(institution_pattern, text, re.IGNORECASE)
        years = re.findall(year_pattern, text, re.IGNORECASE)
        cgpas = re.findall(cgpa_pattern, text, re.IGNORECASE)
        percentages = re.findall(percentage_pattern, text)
        
        # Try to match degree with institution and year
        for i in range(min(len(degrees), len(institutions))):
            edu_entry = {
                "degree": degrees[i].strip(),
                "institution": institutions[i].strip(),
                "year": f"{years[i][0]}-{years[i][1]}" if i < len(years) else "",
                "field": self._extract_field_of_study(text)
            }
            
            # Add CGPA if found
            if cgpas:
                try:
                    edu_entry["cgpa"] = float(cgpas[0][1])
                except:
                    pass
            elif percentages:
                try:
                    edu_entry["percentage"] = float(percentages[0])
                except:
                    pass
            
            education.append(edu_entry)
        
        return education
    
    def _extract_field_of_study(self, text: str) -> str:
        """Extract field of study from common patterns"""
        fields = {
            "Computer Science": ["computer science", "cs", "cse", "information technology", "it"],
            "Electronics": ["electronics", "ece", "electrical"],
            "Mechanical": ["mechanical", "me"],
            "Civil": ["civil", "ce"],
            "Chemical": ["chemical"],
            "Business": ["business", "mba", "management"],
            "Data Science": ["data science", "analytics"]
        }
        
        text_lower = text.lower()
        for field, keywords in fields.items():
            if any(keyword in text_lower for keyword in keywords):
                return field
        
        return "Computer Science"  # Default
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """
        Extract work experience
        
        Looks for:
        - Job titles (Software Engineer, Developer, Intern, etc.)
        - Company names (ending with Pvt Ltd, Inc, Corp, etc.)
        - Duration (Jan 2023 - Dec 2023, June 2025 - Present, etc.)
        - Responsibilities (bullet points)
        """
        experience = []
        
        # Job title patterns
        title_keywords = [
            "Software Engineer", "Developer", "Intern", "Analyst", "Manager",
            "Lead", "Architect", "Consultant", "Designer", "Trainee"
        ]
        title_pattern = r'(' + '|'.join(title_keywords) + r')[\w\s]*'
        
        # Company patterns
        company_pattern = r'[\w\s&]+\s+(Pvt\.?\s*Ltd|Inc\.?|Corp\.?|Corporation|Company|Technologies|Systems|Solutions|Services)'
        
        # Date patterns (Month Year - Month Year)
        date_pattern = r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+(19|20)\d{2}\s*[-–—]\s*((Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+(19|20)\d{2}|Present|present|Current|current)'
        
        # Find all matches
        titles = re.findall(title_pattern, text, re.IGNORECASE)
        companies = re.findall(company_pattern, text, re.IGNORECASE)
        dates = re.findall(date_pattern, text, re.IGNORECASE)
        
        # Match title with company and dates
        for i in range(min(len(titles), len(companies))):
            duration_str = "Not specified"
            if i < len(dates):
                duration_str = f"{dates[i][0]} {dates[i][1]} - {dates[i][3] if dates[i][3] else 'Present'} {dates[i][4] if dates[i][4] else ''}"
            
            exp_entry = {
                "title": titles[i].strip(),
                "company": companies[i].strip(),
                "duration": duration_str,
                "description": self._extract_responsibilities(text, companies[i])
            }
            
            experience.append(exp_entry)
        
        return experience
    
    def _extract_responsibilities(self, text: str, company_name: str) -> List[str]:
        """Extract bullet points near company name"""
        responsibilities = []
        
        # Find section containing company name
        lines = text.split('\n')
        company_index = -1
        
        for i, line in enumerate(lines):
            if company_name.lower() in line.lower():
                company_index = i
                break
        
        if company_index == -1:
            return []
        
        # Extract next 5-10 lines that start with bullet points or dashes
        for i in range(company_index + 1, min(company_index + 10, len(lines))):
            line = lines[i].strip()
            if re.match(r'^[-•·*◦▪]\s+', line) or line.startswith('- '):
                # Clean the bullet point
                clean_line = re.sub(r'^[-•·*◦▪]\s+', '', line)
                if len(clean_line) > 10:  # Avoid too short lines
                    responsibilities.append(clean_line)
        
        return responsibilities[:5]  # Return max 5 responsibilities
    
    def _extract_projects(self, text: str) -> List[str]:
        """Extract project names"""
        projects = []
        
        # Look for "PROJECTS" section
        project_section_pattern = r'PROJECTS?[\s\S]*?(?=EDUCATION|EXPERIENCE|SKILLS|$)'
        match = re.search(project_section_pattern, text, re.IGNORECASE)
        
        if match:
            project_text = match.group(0)
            lines = project_text.split('\n')
            
            for line in lines[1:]:  # Skip header
                line = line.strip()
                # Look for lines that might be project names (capitalized, not too long)
                if (line and 
                    line[0].isupper() and 
                    len(line) < 80 and 
                    not line.lower().startswith('project')):
                    # Clean bullet points
                    clean_line = re.sub(r'^[-•·*◦▪]\s+', '', line)
                    if clean_line:
                        projects.append(clean_line)
        
        return projects[:5]  # Max 5 projects
    
    # ==================== CONFIDENCE SCORING ====================
    
    def _name_confidence(self, name: str, text: str) -> float:
        """Calculate confidence for name extraction"""
        if not name:
            return 0.0
        
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        is_first_line = (lines and lines[0] == name)
        word_count = len(name.split())
        is_title_case = name.istitle()
        
        if is_first_line and 2 <= word_count <= 3 and is_title_case:
            return 0.85
        elif is_first_line:
            return 0.70
        elif 2 <= word_count <= 3:
            return 0.60
        else:
            return 0.40
    
    def _phone_confidence(self, phone: str, text: str) -> float:
        """Calculate confidence for phone extraction"""
        if not phone:
            return 0.10
        
        context_pattern = r'(?:phone|mobile|contact|tel).*?' + re.escape(phone)
        if re.search(context_pattern, text, re.IGNORECASE):
            return 0.90
        elif phone.startswith('+91') or phone.startswith('91'):
            return 0.80
        else:
            return 0.65
    
    def _skills_confidence(self, skills: List[str]) -> float:
        """Calculate confidence based on number of skills found"""
        count = len(skills)
        
        if count == 0:
            return 0.20
        elif count <= 3:
            return 0.55
        elif count <= 8:
            return 0.75
        elif count <= 15:
            return 0.85
        else:
            return 0.90
    
    def _education_confidence(self, education: List[Dict]) -> float:
        """Calculate confidence for education extraction"""
        if not education:
            return 0.20
        
        # Check if we have complete information
        edu = education[0]
        has_degree = bool(edu.get("degree"))
        has_institution = bool(edu.get("institution"))
        has_year = bool(edu.get("year"))
        
        if has_degree and has_institution and has_year:
            return 0.85
        elif has_degree and has_institution:
            return 0.70
        elif has_degree or has_institution:
            return 0.50
        else:
            return 0.30
    
    def _experience_confidence(self, experience: List[Dict]) -> float:
        """Calculate confidence for experience extraction"""
        if not experience:
            return 0.20
        
        exp = experience[0]
        has_title = bool(exp.get("title"))
        has_company = bool(exp.get("company"))
        has_duration = exp.get("duration") != "Not specified"
        has_description = len(exp.get("description", [])) > 0
        
        if has_title and has_company and has_duration and has_description:
            return 0.85
        elif has_title and has_company and has_duration:
            return 0.70
        elif has_title and has_company:
            return 0.55
        else:
            return 0.35
    
    def _empty_result(self) -> Dict:
        """Return empty result structure when parsing fails"""
        return {
            "name": "",
            "email": "",
            "phone": "",
            "skills": [],
            "education": [],
            "experience": [],
            "projects": [],
            "confidence_scores": {
                "name": 0.0,
                "email": 0.0,
                "phone": 0.0,
                "skills": 0.0,
                "education": 0.0,
                "experience": 0.0
            }
        }